"""Estrazione del testo dal curriculum e costruzione del profilo.

Formati supportati: PDF, DOCX, TXT e Markdown. Da questi si ricavano le
competenze, il titolo di studio, le lingue, i ruoli ricoperti e una stima degli
anni di esperienza, che sono gli ingredienti del punteggio di compatibilita'.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import date
from typing import Any

from .skills import (
    education_level_value,
    extract_education_fields,
    extract_education_level,
    extract_languages,
    extract_roles,
    extract_required_years,
    extract_skills,
    resolve_skill,
)
from .text import normalize, normalize_lines


class CVParseError(RuntimeError):
    """Il file non e' leggibile o non contiene testo estraibile."""


# --------------------------------------------------------------------------
# Estrazione del testo
# --------------------------------------------------------------------------

def extract_text(data: bytes, filename: str) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "pdf":
        return _extract_pdf(data)
    if suffix in ("docx", "dotx"):
        return _extract_docx(data)
    if suffix in ("txt", "md", "text", "rtf"):
        return _decode(data)
    if suffix == "doc":
        raise CVParseError(
            "il formato .doc (Word 97-2003) non e' leggibile: converti il file in .docx o in PDF"
        )
    # Ultimo tentativo: se sembra testo, lo si usa comunque.
    text = _decode(data)
    if text.strip():
        return text
    raise CVParseError(f"formato non supportato: .{suffix or 'sconosciuto'}")


def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise CVParseError("libreria pypdf non disponibile") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        # Prima in modalita' "layout": tiene conto della posizione sulla pagina,
        # quindi nei curriculum impaginati a due colonne o dentro tabelle non
        # intreccia le righe. E' quello che rompeva tutto il resto: con le date
        # separate dal loro incarico non si riconosce ne' la sezione ne'
        # l'intervallo. Non tutti i PDF la reggono, e su alcuni restituisce
        # meno testo dell'estrazione semplice: in quel caso si torna indietro.
        pages = _pagine(reader, "layout")
        semplice = _pagine(reader, "")
        if len("".join(semplice).strip()) > len("".join(pages).strip()) * 1.15:
            pages = semplice
    except Exception as exc:
        raise CVParseError(f"PDF illeggibile: {exc}") from exc
    text = "\n".join(pages).strip()
    if not text:
        raise CVParseError(
            "il PDF non contiene testo selezionabile: probabilmente e' una scansione. "
            "Esporta il curriculum in PDF dal documento originale, oppure caricalo in .docx"
        )
    return text


def _pagine(reader: Any, modalita: str) -> list[str]:
    """Testo pagina per pagina, tollerando i PDF che rifiutano una modalita'."""
    fuori = []
    for page in reader.pages:
        try:
            fuori.append(page.extract_text(extraction_mode=modalita) if modalita
                         else page.extract_text() or "")
        except Exception:
            fuori.append("")
    return [p or "" for p in fuori]


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise CVParseError("libreria python-docx non disponibile") from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise CVParseError(f"DOCX illeggibile: {exc}") from exc
    parts = [p.text for p in document.paragraphs]
    # Molti curriculum impaginano le esperienze dentro tabelle.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    text = "\n".join(p for p in parts if p and p.strip()).strip()
    if not text:
        raise CVParseError("il documento Word non contiene testo")
    return text


# --------------------------------------------------------------------------
# Stima degli anni di esperienza
# --------------------------------------------------------------------------

_MONTHS = {
    "gennaio": 1, "gen": 1, "january": 1, "jan": 1,
    "febbraio": 2, "feb": 2, "february": 2,
    "marzo": 3, "mar": 3, "march": 3,
    "aprile": 4, "apr": 4, "april": 4,
    "maggio": 5, "mag": 5, "may": 5,
    "giugno": 6, "giu": 6, "june": 6, "jun": 6,
    "luglio": 7, "lug": 7, "july": 7, "jul": 7,
    "agosto": 8, "ago": 8, "august": 8, "aug": 8,
    "settembre": 9, "set": 9, "sett": 9, "september": 9, "sep": 9, "sept": 9,
    "ottobre": 10, "ott": 10, "october": 10, "oct": 10,
    "novembre": 11, "nov": 11, "november": 11,
    "dicembre": 12, "dic": 12, "december": 12, "dec": 12,
}
_PRESENT = ("oggi", "attuale", "attualmente", "presente", "present", "current", "in corso", "ad oggi")
_MONTH_ALT = "|".join(sorted(_MONTHS, key=len, reverse=True))
_SEP = r"\s*(?:-|--|to|a|al|until|fino a|–|—)\s*"

_RANGE_RE = re.compile(
    # Il guardiano davanti all'anno serve a non leggere "2024" dentro
    # "03/2024": senza, la stessa riga produceva due intervalli - uno da
    # marzo e uno da gennaio - e la fusione teneva il piu' largo,
    # regalando i mesi che stanno prima del mese scritto.
    rf"(?:(?P<m1>{_MONTH_ALT})\s+)?(?<![\d/.])(?P<y1>(?:19|20)\d{{2}})"
    rf"{_SEP}"
    rf"(?:(?:(?P<m2>{_MONTH_ALT})\s+)?(?P<y2>(?:19|20)\d{{2}})|(?P<now>{'|'.join(_PRESENT)}))",
    re.IGNORECASE,
)
_NUMERIC_RANGE_RE = re.compile(
    rf"(?P<m1>\d{{1,2}})[/.](?P<y1>(?:19|20)\d{{2}})"
    rf"{_SEP}"
    rf"(?:(?P<m2>\d{{1,2}})[/.](?P<y2>(?:19|20)\d{{2}})|(?P<now>{'|'.join(_PRESENT)}))",
    re.IGNORECASE,
)

# Letto dall'orologio, non scritto a mano: una data fissa qui dentro fa
# sbagliare tutti gli intervalli aperti ("2019 - oggi") dal primo gennaio
# successivo, e nessuno se ne accorge finche' i conti non tornano piu'.
#
# Si chiede a ogni conteggio e non una volta sola all'importazione: un
# contenitore resta accesso per mesi, e un valore congelato all'avvio sbaglia
# appena passa la mezzanotte di Capodanno.
def _oggi() -> date:
    return date.today()


def _merge_intervals(intervals: list[tuple[int, int]]) -> int:
    """Somma la durata in mesi di intervalli che possono sovrapporsi."""
    if not intervals:
        return 0
    intervals.sort()
    total = 0
    start, end = intervals[0]
    for next_start, next_end in intervals[1:]:
        if next_start <= end:
            end = max(end, next_end)
        else:
            total += end - start
            start, end = next_start, next_end
    total += end - start
    return total


# Intestazioni con cui i curriculum separano le sezioni. Servono a contare gli
# anni di esperienza solo dove sono davvero, senza sommare gli anni di studio.
_SECTION_HEADERS: list[tuple[str, list[str]]] = [
    ("esperienza", [
        "esperienza professionale", "esperienze professionali", "esperienza lavorativa",
        "esperienze lavorative", "esperienza", "work experience", "professional experience",
        "employment history", "employment", "career history", "carriera", "attivita lavorativa",
        "percorso professionale", "experience",
    ]),
    ("formazione", [
        "formazione", "istruzione e formazione", "istruzione", "education",
        "titoli di studio", "percorso formativo", "formazione accademica",
        "academic background", "qualifications", "studi",
    ]),
    ("altro", [
        "competenze", "skills", "capacita e competenze", "lingue", "languages",
        "pubblicazioni", "publications", "certificazioni", "certifications",
        "corsi", "interessi", "hobby", "referenze", "references", "progetti", "projects",
    ]),
]

_ELENCO_INTESTAZIONI = "|".join(
    re.escape(h) for _, hs in _SECTION_HEADERS for h in sorted(hs, key=len, reverse=True))

# L'intestazione da sola sulla sua riga: il caso pulito.
_HEADER_RE = re.compile(rf"^[\s\W]{{0,4}}({_ELENCO_INTESTAZIONI})[\s\W]{{0,4}}$", re.IGNORECASE)
# L'intestazione all'inizio della riga, seguita da altro. Vale solo alle
# condizioni di `_apre_sezione`.
_HEADER_INIZIO_RE = re.compile(rf"^[\s\W]{{0,4}}({_ELENCO_INTESTAZIONI})\b", re.IGNORECASE)

_HEADER_KIND = {h: kind for kind, hs in _SECTION_HEADERS for h in hs}

# Formule che non possono essere l'inizio di una frase normale: se una riga
# comincia cosi' ed e' seguita da altro, e' un titolo di sezione con accanto la
# prima voce. Le parole singole restano fuori di proposito: "Esperienza nella
# gestione di laboratori" e' una riga di requisiti, non un'intestazione.
_INTESTAZIONI_SICURE = frozenset({
    "esperienza professionale", "esperienze professionali", "esperienza lavorativa",
    "esperienze lavorative", "work experience", "professional experience",
    "employment history", "career history", "attivita lavorativa",
    "percorso professionale", "istruzione e formazione", "formazione accademica",
    "titoli di studio", "percorso formativo", "academic background",
    "capacita e competenze",
})


def _apre_sezione(riga: str) -> tuple[str | None, str]:
    """Dice se la riga apre una sezione, e cosa resta della riga dopo il titolo.

    Il caso difficile e' l'intestazione che condivide la riga con la prima
    voce - "ESPERIENZA LAVORATIVA 01/2023 - 06/2024 Tecnico" - che i template
    grafici producono di continuo e che prima non veniva riconosciuta: senza
    sezioni, gli anni di studio finivano sommati a quelli di lavoro.
    """
    pulita = normalize_lines(riga).strip()
    if not pulita:
        return None, ""

    intera = _HEADER_RE.match(pulita)
    if intera:
        return _HEADER_KIND.get(intera.group(1).strip(), "altro"), ""

    inizio = _HEADER_INIZIO_RE.match(pulita)
    if not inizio:
        return None, ""
    titolo = inizio.group(1).strip()
    resto = pulita[inizio.end():].strip(" \t:.-\u2013\u2014|\u00b7\u2022")
    if not resto:
        return _HEADER_KIND.get(titolo, "altro"), ""

    # Due sole vie per accettarla: scritta in maiuscolo nell'originale, oppure
    # una formula inequivocabile.
    lettere = [c for c in riga if c.isalpha()]
    quante = sum(1 for c in titolo if c.isalpha())
    maiuscola = quante > 0 and all(c.isupper() for c in lettere[:quante])
    if maiuscola or titolo in _INTESTAZIONI_SICURE:
        return _HEADER_KIND.get(titolo, "altro"), resto
    return None, ""


def split_sections(text: str) -> list[tuple[str, str]]:
    """Divide il curriculum in sezioni (tipo, contenuto) usando le intestazioni.

    Se non riconosce nessuna intestazione restituisce il testo come sezione
    unica di tipo sconosciuto, e chi chiama decide come comportarsi.
    """
    sezioni: list[tuple[str, str]] = []
    tipo = "intestazione"
    corpo: list[str] = []
    trovata = False

    for riga in text.splitlines():
        apre, resto = _apre_sezione(riga)
        if apre is None:
            corpo.append(normalize_lines(riga))
            continue
        trovata = True
        sezioni.append((tipo, "\n".join(corpo).strip()))
        tipo, corpo = apre, ([resto] if resto else [])
    sezioni.append((tipo, "\n".join(corpo).strip()))

    if not trovata:
        return [("sconosciuto", normalize_lines(text))]
    return [(k, b) for k, b in sezioni if b or k != "intestazione"]


def estimate_years(text: str) -> float:
    """Stima gli anni di esperienza sommando gli intervalli di date del CV.

    Conta solo le date che cadono nelle sezioni di esperienza lavorativa: senza
    questo filtro gli anni di universita' verrebbero sommati a quelli di lavoro
    e ogni candidato risulterebbe molto piu' esperto di quanto sia.

    Gli intervalli sovrapposti (lavori paralleli, studio piu' lavoro) vengono
    fusi, altrimenti un curriculum ricco gonfierebbe comunque il totale.
    """
    sections = split_sections(text)
    work = [body for kind, body in sections if kind in ("esperienza", "sconosciuto")]
    # Nessuna sezione di esperienza riconosciuta: si ripiega sull'intero testo,
    # accettando una stima piu' grossolana.
    normalized = "\n".join(work) if work else normalize(text)
    intervals: list[tuple[int, int]] = []
    oggi = _oggi()

    for regex, numeric in ((_RANGE_RE, False), (_NUMERIC_RANGE_RE, True)):
        for match in regex.finditer(normalized):
            year1 = int(match.group("y1"))
            if numeric:
                month1 = int(match.group("m1") or 1)
            else:
                month1 = _MONTHS.get((match.group("m1") or "").lower(), 1)
            if match.group("now"):
                # Un lavoro in corso finisce adesso, non alla fine dell'anno.
                # Chiuderlo a dicembre regalava i mesi non ancora vissuti: a
                # settembre "2019 - oggi" contava tre mesi di troppo, e
                # "gennaio 2026 - oggi" ne contava dodici invece di nove.
                year2, month2 = oggi.year, oggi.month
            else:
                year2 = int(match.group("y2"))
                month2 = int(match.group("m2") or 12) if numeric else _MONTHS.get((match.group("m2") or "").lower(), 12)
            if not (1950 <= year1 <= oggi.year and year1 <= year2 <= oggi.year + 1):
                continue
            start = year1 * 12 + max(1, min(12, month1))
            end = year2 * 12 + max(1, min(12, month2))
            if 0 < end - start <= 12 * 45:
                intervals.append((start, end))

    months = _merge_intervals(intervals)
    if months:
        return round(months / 12.0, 1)
    # Nessun intervallo: molti curriculum scrivono un anno solo e accanto la
    # durata ("2023 Tirocinio in laboratorio (6 mesi)"). Prima non contava
    # niente e il candidato risultava senza esperienza.
    durate = _durate_dichiarate(normalized)
    if durate:
        return round(durate / 12.0, 1)
    # Ultima possibilita': la dichiarazione esplicita, "3 anni di esperienza".
    declared = extract_required_years(text)
    return float(declared) if declared else 0.0


# "(6 mesi)", "18 mesi", "2 anni": la durata scritta a parole accanto a una voce.
_DURATA_VOCE = re.compile(r"\b(\d{1,2})\s*(mesi|mese|anni|anno)\b", re.IGNORECASE)
_HA_ANNO = re.compile(r"\b(?:19|20)\d{2}\b")


def _durate_dichiarate(testo: str) -> int:
    """Somma in mesi le durate scritte accanto a una voce datata.

    Solo righe che portano anche un anno: senza quel vincolo si sommerebbe
    anche "cerco un impiego di almeno 2 anni" dalla lettera di presentazione.
    """
    totale = 0
    for riga in testo.splitlines():
        if not _HA_ANNO.search(riga):
            continue
        for match in _DURATA_VOCE.finditer(riga):
            n = int(match.group(1))
            mesi = n if match.group(2).startswith("mes") else n * 12
            if 0 < mesi <= 12 * 45:
                totale += mesi
    return totale


# --------------------------------------------------------------------------
# Nome della persona
# --------------------------------------------------------------------------

# Righe che non possono essere un nome: contatti, date, elenchi puntati.
_NON_NOME = re.compile(r"[0-9@|/\\•·+_=]|https?:|www\.", re.IGNORECASE)

# Parole che compaiono in testa a un curriculum ma non sono nomi di persona.
_PAROLE_NON_NOME = {
    "curriculum", "vitae", "cv", "resume", "resumé", "profilo", "profile",
    "contatti", "contatto", "contact", "contacts", "lingue", "languages",
    "esperienza", "esperienze", "experience", "formazione", "education",
    "competenze", "skills", "istruzione", "studi", "obiettivo", "about",
    "informazioni", "personali", "personal", "dati", "riepilogo", "summary",
    "dottore", "dottoressa", "ingegnere", "avvocato", "sig", "sig.ra",
    "europass", "telefono", "cellulare", "indirizzo", "email", "mail",
    "nazionalita", "residenza", "domicilio", "portfolio", "linkedin",
}


def _sembra_nome(riga: str) -> list[str] | None:
    """Dice se una riga puo' essere un nome di persona, e la scompone in parole."""
    if not riga or len(riga) > 42 or _NON_NOME.search(riga):
        return None
    parole = riga.replace(",", " ").split()
    if not 1 <= len(parole) <= 4:
        return None
    for p in parole:
        pulita = p.replace("'", "").replace("-", "").replace(".", "")
        if not pulita.isalpha() or not 2 <= len(pulita) <= 20:
            return None
        if not p[0].isupper():
            return None
        if pulita.lower() in _PAROLE_NON_NOME:
            return None
    # Una riga tutta maiuscola con tre o piu' parole e' quasi sempre una
    # qualifica ("DOTTORE IN SCIENZE BIOLOGICHE"). Con due parole invece puo'
    # benissimo essere un nome scritto in maiuscolo ("LUCA VERDI").
    if riga.isupper() and len(parole) > 2:
        return None
    return parole


def extract_person_name(text: str) -> str:
    """Ricava nome e cognome dalle prime righe del curriculum.

    Due casi da coprire: il nome su una riga sola ("Mario Rossi") e il nome
    spezzato su righe consecutive, che capita spesso perche' nei PDF impaginati
    a colonne nome e cognome finiscono su capoversi diversi.

    Restituisce stringa vuota se non trova nulla di credibile: meglio ricadere
    sul nome del file che inventare un nome sbagliato.
    """
    righe = [r.strip() for r in text.splitlines()]
    righe = [r for r in righe if r][:12]
    if not righe:
        return ""

    # Si guarda solo la sequenza iniziale: il nome sta in cima, e fermarsi al
    # primo capoverso che non somiglia a un nome evita di pescare piu' avanti
    # parole isolate come "Lingue" o "Roma".
    sequenza: list[list[str]] = []
    for riga in righe:
        parole = _sembra_nome(riga)
        if parole is None:
            # Finche' non si e' trovato nulla si tira dritto: in cima possono
            # esserci "Curriculum Vitae" o un recapito. Una volta iniziata la
            # sequenza, invece, la prima riga diversa la chiude.
            if sequenza:
                break
            continue
        sequenza.append(parole)

    if not sequenza:
        return ""
    if len(sequenza[0]) >= 2:
        parole = sequenza[0][:4]
    else:
        # Nome e cognome su righe separate. Ci si ferma a due: una terza riga
        # da una parola sola e' piu' spesso la qualifica ("Farmacista") che un
        # secondo cognome, e sbagliare qui mette una professione al posto del nome.
        parole = [p[0] for p in sequenza[:2] if len(p) == 1]
        if len(parole) < 2:
            return ""

    nome = " ".join(parole)
    return nome.title() if nome.isupper() else nome


# --------------------------------------------------------------------------
# Profilo
# --------------------------------------------------------------------------

@dataclass
class CVProfile:
    raw_text: str
    skills: list[str] = field(default_factory=list)
    education_fields: list[str] = field(default_factory=list)
    education_level: int = 0
    education_label: str = ""
    languages: list[str] = field(default_factory=list)
    roles: list[str] = field(default_factory=list)
    years_experience: float = 0.0

    def to_storage(self) -> dict[str, Any]:
        return {
            "skills": self.skills,
            "education": {
                "level": self.education_level,
                "label": self.education_label,
                "fields": self.education_fields,
            },
            "titles": self.roles,
            "languages": self.languages,
            "years_experience": self.years_experience,
        }


# Scarto oltre il quale due letture degli anni di esperienza non si possono
# considerare d'accordo. Sotto l'anno e' arrotondamento; sopra, uno dei due ha
# letto il curriculum in un modo che l'altro non riconosce.
SCARTO_ANNI_TOLLERATO = 1.0


def apply_reading(profile: CVProfile, reading: Any) -> dict[str, Any]:
    """Fonde nel profilo la lettura del modello, e riporta i disaccordi.

    Le euristiche restano la base: sono verificabili riga per riga e funzionano
    senza chiavi. Il modello aggiunge quello che il dizionario non copre e
    corregge gli anni dove le date da sole non bastavano.

    Dove le due letture non vanno d'accordo il disaccordo viene registrato
    invece che risolto in silenzio: e' l'unico modo che ha chi guarda il
    profilo di accorgersi che il modello ha preso un abbaglio.
    """
    note: dict[str, Any] = {"fonte": "modello", "divergenze": [], "esperienze": [],
                            "extra_tags": [], "avviso": ""}
    if reading is None:
        note["fonte"] = "euristiche"
        return note

    # -- competenze: quelle note entrano nel confronto, le altre restano
    #    etichette libere, che e' comunque meglio di buttarle via.
    extra: list[str] = []
    for grezza in getattr(reading, "skills", []) or []:
        pulita = (grezza or "").strip()[:60]
        if not pulita:
            continue
        canonica = resolve_skill(pulita)
        if canonica:
            if canonica not in profile.skills:
                profile.skills.append(canonica)
        elif pulita.lower() not in {e.lower() for e in extra}:
            extra.append(pulita)
    note["extra_tags"] = extra

    for grezza in getattr(reading, "languages", []) or []:
        for lingua in extract_languages(grezza or ""):
            if lingua not in profile.languages:
                profile.languages.append(lingua)

    for grezza in getattr(reading, "roles", []) or []:
        for ruolo in extract_roles(grezza or ""):
            if ruolo not in profile.roles:
                profile.roles.append(ruolo)

    for grezza in getattr(reading, "education_fields", []) or []:
        for campo in extract_education_fields(grezza or ""):
            if campo not in profile.education_fields:
                profile.education_fields.append(campo)

    # -- titolo di studio
    etichetta = (getattr(reading, "education_level", "") or "").strip()
    livello = education_level_value(etichetta)
    if livello:
        if profile.education_level and profile.education_level != livello:
            note["divergenze"].append(
                f"titolo di studio: le date e le parole del curriculum dicono "
                f"«{profile.education_label}», il modello legge «{etichetta}»")
        profile.education_level, profile.education_label = livello, etichetta

    # -- anni di esperienza
    anni = float(getattr(reading, "years_experience", 0) or 0)
    if anni > 0 and profile.years_experience > 0 and \
            abs(anni - profile.years_experience) >= SCARTO_ANNI_TOLLERATO:
        note["divergenze"].append(
            f"anni di esperienza: dalle date risultano {profile.years_experience:g}, "
            f"il modello ne conta {anni:g}")
        profile.years_experience = anni
    elif anni > 0 and profile.years_experience <= 0:
        profile.years_experience = anni
    elif anni <= 0 and profile.years_experience > 0:
        note["divergenze"].append(
            f"anni di esperienza: dalle date risultano {profile.years_experience:g}, "
            "il modello non ne ha riconosciuto nessuno - tenuto il conto delle date")

    note["esperienze"] = [str(x)[:160] for x in (getattr(reading, "experience_items", []) or [])[:10]]
    note["avviso"] = (getattr(reading, "notes", "") or "")[:300]
    return note


def build_profile(text: str) -> CVProfile:
    """Ricava dal testo del curriculum tutto cio' che serve al punteggio."""
    # Il titolo di studio si cerca dentro la sezione che lo riguarda, non in
    # tutto il documento: un dottorato citato fra le pubblicazioni o nel
    # racconto di un incarico faceva risultare dottore chi ha la maturita'.
    # Senza sezioni riconosciute si ripiega sul testo intero, com'era prima.
    formazione = "\n".join(b for k, b in split_sections(text) if k == "formazione")
    level, label = extract_education_level(formazione or text)
    return CVProfile(
        raw_text=text,
        skills=extract_skills(text),
        education_fields=extract_education_fields(text),
        education_level=level,
        education_label=label,
        languages=extract_languages(text),
        roles=extract_roles(text),
        years_experience=estimate_years(text),
    )
